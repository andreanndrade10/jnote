import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jnote import *

def write_table_rede_cabeada(school_name):
    index = df_questions[df_questions['Unidade']==school_name].index.tolist()[0]

    records = (
        ('Fabricante',df_questions[14][index],df_questions[23][index],df_questions[32][index],df_questions[43][index]),
        ('Modelo',df_questions[15][index],df_questions[24][index],df_questions[33][index],df_questions[44][index]),
        ('Total de portas',df_questions[19][index],df_questions[28][index],df_questions[37][index],'---'),
        ('Total de SW do modelo',df_questions[17][index],df_questions[26][index],df_questions[35][index],df_questions[45][index])
    )

    table = document.add_table(rows=1, cols=5, style='Medium Shading 1 Accent 5')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '        '
    hdr_cells[1].text = 'Modelo 3'
    hdr_cells[2].text = 'Modelo 2'
    hdr_cells[3].text = 'Modelo 1'
    hdr_cells[4].text = 'Switch L3'

    for a,b,c,d,e in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(a)
        row_cells[1].text = str(b)
        row_cells[2].text = str(c)
        row_cells[3].text = str(d)
        row_cells[4].text = str(e)

def write_table_wireless(school_name):
    index = df_questions[df_questions['Unidade']==school_name].index.tolist()[0]
    records = (
        ('Fabricante',df_questions[64][index], df_questions[71][index]),
        ('Modelo', df_questions[65][index], df_questions[72][index]),
        ('MU-MIMO', df_questions[70][index], '---'),
        ('Quantidade de APs Suportados','---',df_questions[73][index]),
        ('Quantidade de Clientes total suportada','---',df_questions[74][index])
    )

    table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 5')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '        '
    hdr_cells[1].text = 'Access Point'
    hdr_cells[2].text = 'Controladora'

    for a,b,c in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(a)
        row_cells[1].text = str(b)
        row_cells[2].text = str(c)
 
def write_table_conectividade(school_name):
    index = df_questions[df_questions['Unidade']==school_name].index.tolist()[0]
    records = (
        ('Banda do Link',df_questions[79][index], df_questions[84][index]),
        ('Tipo de Link', df_questions[78][index], df_questions[83][index]),
        ('Provedor', df_questions[80][index], df_questions[85][index])
    )

    table = document.add_table(rows=1, cols=3, style='Medium Shading 1 Accent 5')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '        '
    hdr_cells[1].text = 'Link Primário (Mbps)'
    hdr_cells[2].text = 'Link Secundário (Mbps)'

    for a,b,c in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(a)
        row_cells[1].text = str(b)
        row_cells[2].text = str(c)

def write_table_controle_de_conteudo(school_name):
    index = df_questions[df_questions['Unidade']==school_name].index.tolist()[0]
    records = (
        ('Modelo', df_questions[90][index], df_questions[97][index]),
        ('Fabricante', df_questions[89][index], '---'),
        ('Tipo', df_questions[88][index], '---')
    )

    table = document.add_table(rows=1, cols=3 , style='Medium Shading 1 Accent 5')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '        '
    hdr_cells[1].text = 'Firewall'
    hdr_cells[2].text = 'NAC'

    for a,b,c in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(a)
        row_cells[1].text = str(b)
        row_cells[2].text = str(c)
############
def write_unit_tables(school_name):

    center=document.add_paragraph('Rede Cabeada')
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_table_rede_cabeada(school_name)

    center=document.add_paragraph('Wireless')
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_table_wireless(school_name)

    center=document.add_paragraph('Conectividade')
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_table_conectividade(school_name)

    center=document.add_paragraph('Controle de Conteúdo')
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_table_controle_de_conteudo(school_name)

def write_specifics(school_name):
    document.add_paragraph(specific_infra(school_name))
    document.add_paragraph(specific_connection(school_name))
    document.add_paragraph(specific_contentMaturity(school_name))


def write_plot_class(lugar):
    plot_graph_class(lugar)
    img=document.add_picture('img/graph_class/graph_class_{}.png'.format(lugar), width=Inches(5.0))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def write_plot_size(lugar):
    plot_graph_size(lugar)
    img=document.add_picture('img/graph_size/graph_size_{}.png'.format(lugar), width=Inches(5.0))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def write_regiao(regiao):
    document.add_heading(regiao, level=1)
    write_plot_class(regiao)
    #write_percentage_class(regiao)
    write_plot_size(regiao)
    #write_percentage_size(regiao)

def write_estado(estado):
    document.add_heading(estado, level=2)
    write_plot_class(estado)
    #write_percentage_class(estado)
    write_plot_size(estado)
    #write_percentage_size(estado)

def write_unidades(unidade):
    a = document.add_paragraph('Classe pertencente: ')
    a.add_run(get_classe(unidade))
    document.add_heading(unidade, level=3)
    write_unit_tables(unidade)
    p = document.add_paragraph('Maior prioridade de mudança: ')
    p.add_run(change_priority(unidade)).italic = True
    write_specifics(unidade)

def write_doc():
    for i in df['Regiao'].unique():
        write_regiao(i)
        document.add_page_break()
        for j in df[df.Regiao==i]['UF'].unique():
            write_estado(j)
            document.add_page_break()
            for k in df[df.UF==j]['Unidade']:
                write_unidades(k)
                document.add_page_break()
                


if __name__ == "__main__":
    document = Document()
    write_doc()
    document.save('escolas.docx')